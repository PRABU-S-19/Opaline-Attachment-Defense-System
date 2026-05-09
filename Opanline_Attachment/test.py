import docx

doc = docx.Document("static/attachments/f1_doc2.docx")

all_paras = doc.paragraphs
s=len(all_paras)
for para in all_paras:
    print(para.text)
    print("-------")





'''from docx import Document
from io import StringIO
#f = open('static/attachments/f1_doc2.docx', 'rb')
#document = Document(f)
#f.close()

# or

with open('static/attachments/f1_doc2.docx', 'rb') as f:
    source_stream = StringIO(f.read())
document = Document(source_stream)
print(document)'''
